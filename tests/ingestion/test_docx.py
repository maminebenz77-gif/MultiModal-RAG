import io
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

from multimodal_rag.ingestion.docx import parse_docx
from multimodal_rag.ingestion.schema import ElementType
from multimodal_rag.providers.base import VisionProvider


class FakeVisionProvider(VisionProvider):
    def describe(self, image_bytes: bytes, prompt: str | None = None) -> str:
        return f"a description of {len(image_bytes)} bytes"


@pytest.fixture(autouse=True)
def _fake_vision(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(
        "multimodal_rag.ingestion.vision.get_vision", lambda: FakeVisionProvider()
    )


def _write_sample(tmp_path: Path, image_bytes: bytes) -> Path:
    image_path = tmp_path / "pic.png"
    image_path.write_bytes(image_bytes)

    document = Document()
    document.add_heading("A Title", level=1)
    document.add_paragraph("A paragraph of body text.")

    table = document.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "A", "B"
    row = table.add_row().cells
    row[0].text, row[1].text = "1", "2"

    document.add_picture(str(image_path))
    document.add_paragraph("A closing paragraph.")

    docx_path = tmp_path / "doc.docx"
    document.save(str(docx_path))
    return docx_path


def _tiny_png() -> bytes:
    # A real, valid 1x1 PNG — python-docx needs to parse an actual PNG
    # header (via Pillow's encoder here) to determine image dimensions.
    buffer = io.BytesIO()
    Image.new("RGB", (1, 1), color="red").save(buffer, format="PNG")
    return buffer.getvalue()


_TINY_PNG = _tiny_png()


def test_parses_title_paragraph_table_and_image(tmp_path: Path) -> None:
    elements = parse_docx(_write_sample(tmp_path, _TINY_PNG))

    types = [el.type for el in elements]
    assert types == [
        ElementType.TITLE,
        ElementType.PARAGRAPH,
        ElementType.TABLE,
        ElementType.IMAGE,
        ElementType.PARAGRAPH,
    ]


def test_title_uses_heading_style(tmp_path: Path) -> None:
    elements = parse_docx(_write_sample(tmp_path, _TINY_PNG))
    assert elements[0].text == "A Title"


def test_table_rendered_as_markdown(tmp_path: Path) -> None:
    elements = parse_docx(_write_sample(tmp_path, _TINY_PNG))
    table = next(el for el in elements if el.type == ElementType.TABLE)
    assert table.text == "| A | B |\n| --- | --- |\n| 1 | 2 |"


def test_image_element_has_bytes_and_description(tmp_path: Path) -> None:
    elements = parse_docx(_write_sample(tmp_path, _TINY_PNG))
    image = next(el for el in elements if el.type == ElementType.IMAGE)
    assert image.image_bytes == _TINY_PNG
    assert image.description == f"a description of {len(_TINY_PNG)} bytes"
    assert image.description_status == "generated"


def test_document_order_is_preserved(tmp_path: Path) -> None:
    elements = parse_docx(_write_sample(tmp_path, _TINY_PNG))
    assert [el.metadata.position for el in elements] == [0, 1, 2, 3, 4]
