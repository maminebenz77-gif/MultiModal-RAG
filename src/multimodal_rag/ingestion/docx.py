"""DOCX parser: python-docx -> common Element schema.

Walks the document body in document order — python-docx's own
`document.paragraphs`/`document.tables` are separate flat lists that lose
interleaving, so a table or image in the middle of the text would end up
in the wrong position relative to surrounding paragraphs.

DOCX gives paragraph *style names* directly ("Heading 1", "Title", ...),
which is a reliable title signal — no font-size heuristic needed, unlike
PDF.

Known gap: a native Word "Chart" object (an embedded, editable chart, as
opposed to a pasted-in raster image of one) uses different XML (a
graphicFrame referencing chart data) than a picture blip, and isn't
detected here — only raster images are. Same as PDF/Markdown, an
embedded picture *of* a chart is just an "image" element; there's no way
to tell it's a chart without a vision description.
"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from .schema import Element, ElementMetadata, ElementType
from .tables import rows_to_markdown, summarize_table
from .vision import ImageDescriber

_TITLE_STYLES = {
    "Title",
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Heading 4",
    "Heading 5",
    "Heading 6",
}


def parse_docx(path: Path, summarize_tables: bool = False) -> list[Element]:
    document = Document(str(path))
    describer = ImageDescriber()

    elements: list[Element] = []
    position = 0

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = DocxParagraph(child, document)
            image_rids = _embedded_image_rids(paragraph)
            if image_rids:
                for rid in image_rids:
                    image_bytes = document.part.related_parts[rid].blob
                    description, status = describer.describe(image_bytes)
                    elements.append(
                        Element(
                            type=ElementType.IMAGE,
                            image_bytes=image_bytes,
                            description=description,
                            description_status=status,
                            metadata=ElementMetadata(source_file=str(path), position=position),
                        )
                    )
                    position += 1
                continue

            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else None
            element_type = (
                ElementType.TITLE if style_name in _TITLE_STYLES else ElementType.PARAGRAPH
            )
            elements.append(
                Element(
                    type=element_type,
                    text=text,
                    metadata=ElementMetadata(source_file=str(path), position=position),
                )
            )
            position += 1

        elif child.tag == qn("w:tbl"):
            table = DocxTable(child, document)
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            markdown_table = rows_to_markdown(rows)
            summary = summarize_table(markdown_table) if summarize_tables else None
            elements.append(
                Element(
                    type=ElementType.TABLE,
                    text=markdown_table,
                    table_summary=summary,
                    metadata=ElementMetadata(source_file=str(path), position=position),
                )
            )
            position += 1

    return elements


def _embedded_image_rids(paragraph: DocxParagraph) -> list[str]:
    """Relationship IDs of any inline raster images embedded in this paragraph."""
    blips = paragraph._element.findall(".//" + qn("a:blip"))
    rids = [blip.get(qn("r:embed")) for blip in blips]
    return [rid for rid in rids if rid]
